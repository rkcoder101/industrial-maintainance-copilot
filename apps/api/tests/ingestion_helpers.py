from io import BytesIO
from pathlib import Path
from zipfile import ZipFile

import fitz  # type: ignore[import-untyped]
from docx import Document as DocxDocument
from fastapi import UploadFile
from openpyxl import Workbook  # type: ignore[import-untyped]
from PIL import Image, ImageDraw


def upload_file(name: str, content: bytes) -> UploadFile:
    return UploadFile(BytesIO(content), filename=name)


def pdf_bytes() -> bytes:
    return b"%PDF-1.4\n1 0 obj\n<<>>\nendobj\n%%EOF\n"


def valid_pdf_bytes(pages: list[str] | None = None) -> bytes:
    document = fitz.open()
    for text in pages or ["P-101 Pump Maintenance\nInspect seal pressure daily."]:
        page = document.new_page()
        page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


def png_bytes() -> bytes:
    return b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR"


def valid_png_bytes(text: str = "P-101 OK") -> bytes:
    image = Image.new("RGB", (360, 120), "white")
    draw = ImageDraw.Draw(image)
    draw.text((20, 45), text, fill="black")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def jpeg_bytes() -> bytes:
    return b"\xff\xd8\xff\xe0JFIF\x00\xff\xd9"


def csv_bytes() -> bytes:
    return b"asset,value\nP-101,42\n"


def docx_bytes() -> bytes:
    return _office_zip({"[Content_Types].xml": b"<Types/>", "word/document.xml": b"<doc/>"})


def valid_docx_bytes() -> bytes:
    document = DocxDocument()
    document.add_heading("Safety Procedure", level=1)
    document.add_paragraph("Lock out P-101 before inspection.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Asset"
    table.cell(0, 1).text = "Action"
    table.cell(1, 0).text = "P-101"
    table.cell(1, 1).text = "Inspect bearing"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def xlsx_bytes() -> bytes:
    return _office_zip({"[Content_Types].xml": b"<Types/>", "xl/workbook.xml": b"<workbook/>"})


def valid_xlsx_bytes() -> bytes:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Inspection"
    sheet.append(["Asset", "Reading"])
    sheet.append(["P-101", "42"])
    buffer = BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def invalid_office_zip_bytes() -> bytes:
    return _office_zip({"[Content_Types].xml": b"<Types/>", "other/file.xml": b"<x/>"})


def _office_zip(files: dict[str, bytes]) -> bytes:
    buffer = BytesIO()
    with ZipFile(buffer, "w") as archive:
        for name, content in files.items():
            archive.writestr(name, content)
    return buffer.getvalue()


def assert_inside(path: Path, root: Path) -> None:
    assert path.resolve().is_relative_to(root.resolve())
