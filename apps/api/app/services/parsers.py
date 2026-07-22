import csv
import os
import re
import shutil
import tempfile
from abc import ABC, abstractmethod
from collections.abc import Iterable
from io import StringIO
from pathlib import Path
from typing import Protocol

import fitz  # type: ignore[import-untyped]
import pytesseract  # type: ignore[import-untyped]
from charset_normalizer import from_bytes
from docx import Document as DocxDocument
from openpyxl import load_workbook  # type: ignore[import-untyped]
from PIL import Image, UnidentifiedImageError

from app.models.enums import BlockType
from app.services.parsing_contracts import (
    PARSER_VERSION,
    ParseContext,
    ParsedBlockData,
    ParsedDocumentData,
    ParsedPageData,
    ParserWarning,
)
from app.services.processing_errors import (
    DocumentParseError,
    PageRenderError,
    UnsupportedParserError,
)
from app.services.storage import LocalStorageBackend

TEXT_FILE_TYPES = {"txt", "log"}
CSV_FILE_TYPES = {"csv"}
IMAGE_FILE_TYPES = {"png", "jpg", "jpeg"}
DOCX_FILE_TYPES = {"docx"}
XLSX_FILE_TYPES = {"xlsx"}
PDF_FILE_TYPES = {"pdf"}


class DocumentParser(Protocol):
    name: str
    parser_version: str

    def supports(self, file_type: str | None, mime_type: str | None) -> bool: ...

    def parse(self, context: ParseContext) -> ParsedDocumentData: ...


class BaseParser(ABC):
    name: str
    parser_version = PARSER_VERSION

    @abstractmethod
    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        raise NotImplementedError

    @abstractmethod
    def parse(self, context: ParseContext) -> ParsedDocumentData:
        raise NotImplementedError

    def _document(
        self,
        *,
        pages: list[ParsedPageData],
        warnings: list[ParserWarning] | None = None,
        metadata: dict[str, object] | None = None,
        fallback_parser_name: str | None = None,
        ocr_used: bool = False,
    ) -> ParsedDocumentData:
        return ParsedDocumentData(
            parser_name=self.name,
            parser_version=self.parser_version,
            pages=pages,
            warnings=warnings or [],
            metadata=metadata or {},
            fallback_parser_name=fallback_parser_name,
            ocr_used=ocr_used,
        )


class PageImageWriter:
    def write_pymupdf_page(self, context: ParseContext, page: fitz.Page, page_number: int) -> str:
        key = context.rendered_storage.make_page_image_key(
            document_id=context.document_id,
            page_number=page_number,
        )
        temp_path = self._make_temp_path(context, page_number)
        try:
            zoom = context.settings.page_render_dpi / 72
            pixmap = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            pixmap.save(temp_path)
            self._move_to_storage(context, temp_path, key)
        except Exception as exc:
            temp_path.unlink(missing_ok=True)
            raise PageRenderError() from exc
        return key

    def write_image(self, context: ParseContext, image: Image.Image, page_number: int) -> str:
        key = context.rendered_storage.make_page_image_key(
            document_id=context.document_id,
            page_number=page_number,
        )
        temp_path = self._make_temp_path(context, page_number)
        try:
            image.convert("RGB").save(temp_path, format="PNG")
            self._move_to_storage(context, temp_path, key)
        except Exception as exc:
            temp_path.unlink(missing_ok=True)
            raise PageRenderError() from exc
        return key

    def _make_temp_path(self, context: ParseContext, page_number: int) -> Path:
        context.settings.rendered_pages_temp_path.mkdir(parents=True, exist_ok=True, mode=0o750)
        handle, name = tempfile.mkstemp(
            prefix=f"{context.document_id}-p{page_number}-",
            suffix=".png",
            dir=context.settings.rendered_pages_temp_path,
        )
        os.close(handle)
        return Path(name)

    def _move_to_storage(self, context: ParseContext, temp_path: Path, key: str) -> None:
        backend = context.rendered_storage.backend
        if not isinstance(backend, LocalStorageBackend):
            raise PageRenderError("Only local rendered page storage is supported in Phase 4.")
        destination = backend.resolve_key(key)
        destination.parent.mkdir(parents=True, exist_ok=True, mode=0o750)
        os.replace(temp_path, destination)


class PyMuPDFParser(BaseParser):
    name = "pymupdf"

    def __init__(self, image_writer: PageImageWriter | None = None) -> None:
        self.image_writer = image_writer or PageImageWriter()

    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        return (file_type or "").lower() in PDF_FILE_TYPES or mime_type == "application/pdf"

    def parse(self, context: ParseContext) -> ParsedDocumentData:
        warnings = [
            ParserWarning(
                code="docling_unavailable",
                message="Docling parser is not installed; PyMuPDF parsed the PDF.",
                scope="document",
            )
        ]
        try:
            document = fitz.open(context.file_path)
        except Exception as exc:
            raise DocumentParseError("PDF parsing failed.") from exc

        pages: list[ParsedPageData] = []
        try:
            for index, page in enumerate(document, start=1):
                text = _normalize_text(page.get_text("text"))
                block_entries = page.get_text("blocks")
                blocks = _blocks_from_pymupdf(block_entries, page_number=index)
                image_key = self.image_writer.write_pymupdf_page(context, page, index)
                pages.append(
                    ParsedPageData(
                        page_number=index,
                        logical_page_key=f"page-{index}",
                        text=text,
                        blocks=blocks,
                        rendered_image_key=image_key,
                        width=float(page.rect.width),
                        height=float(page.rect.height),
                        parser_metadata={"parser": self.name},
                    )
                )
        finally:
            document.close()

        return self._document(
            pages=pages,
            warnings=warnings,
            metadata={"source_format": "pdf"},
            fallback_parser_name="docling",
        )


class DocxParser(BaseParser):
    name = "python-docx"

    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        return (file_type or "").lower() in DOCX_FILE_TYPES or mime_type == (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    def parse(self, context: ParseContext) -> ParsedDocumentData:
        try:
            document = DocxDocument(str(context.file_path))
        except Exception as exc:
            raise DocumentParseError("DOCX parsing failed.") from exc

        blocks: list[ParsedBlockData] = []
        section_path: list[str] = []
        for paragraph in document.paragraphs:
            text = _normalize_text(paragraph.text)
            if not text:
                continue
            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            heading_level = _heading_level(style_name)
            block_type = BlockType.HEADING if heading_level else BlockType.PARAGRAPH
            if heading_level:
                section_path = section_path[: heading_level - 1]
                section_path.append(text)
            blocks.append(
                ParsedBlockData(
                    block_index=len(blocks),
                    block_type=block_type,
                    text=text,
                    page_number=1,
                    heading_level=heading_level,
                    section_path=" > ".join(section_path) if section_path else None,
                    parser_metadata={"style": paragraph.style.name if paragraph.style else None},
                )
            )

        for table_index, table in enumerate(document.tables, start=1):
            rows = [
                [_normalize_text(cell.text) for cell in row.cells]
                for row in table.rows
                if any(_normalize_text(cell.text) for cell in row.cells)
            ]
            if not rows:
                continue
            text = _rows_to_markdown(rows)
            blocks.append(
                ParsedBlockData(
                    block_index=len(blocks),
                    block_type=BlockType.TABLE,
                    text=text,
                    page_number=1,
                    section_path=" > ".join(section_path) if section_path else None,
                    table_metadata={"row_count": len(rows), "table_index": table_index},
                )
            )

        page_text = "\n\n".join(block.text for block in blocks)
        page = ParsedPageData(
            page_number=1,
            logical_page_key="document",
            text=page_text,
            blocks=blocks,
            parser_metadata={"parser": self.name},
        )
        return self._document(pages=[page], metadata={"source_format": "docx"})


class OpenPyXLParser(BaseParser):
    name = "openpyxl"

    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        return (file_type or "").lower() in XLSX_FILE_TYPES or mime_type == (
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    def parse(self, context: ParseContext) -> ParsedDocumentData:
        try:
            workbook = load_workbook(context.file_path, read_only=True, data_only=True)
        except Exception as exc:
            raise DocumentParseError("XLSX parsing failed.") from exc

        pages: list[ParsedPageData] = []
        for page_number, sheet in enumerate(workbook.worksheets, start=1):
            rows: list[list[str]] = []
            for row in sheet.iter_rows(values_only=True):
                cells = [_normalize_text("" if cell is None else str(cell)) for cell in row]
                if any(cells):
                    rows.append(cells)
            text = _rows_to_markdown(rows)
            blocks = (
                [
                    ParsedBlockData(
                        block_index=0,
                        block_type=BlockType.TABLE,
                        text=text,
                        page_number=page_number,
                        section_path=sheet.title,
                        table_metadata={"row_count": len(rows), "sheet_name": sheet.title},
                    )
                ]
                if text
                else []
            )
            pages.append(
                ParsedPageData(
                    page_number=page_number,
                    logical_page_key=sheet.title,
                    text=text,
                    blocks=blocks,
                    parser_metadata={"sheet_name": sheet.title, "parser": self.name},
                )
            )
        workbook.close()
        return self._document(pages=pages, metadata={"source_format": "xlsx"})


class CSVParser(BaseParser):
    name = "csv"

    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        return (file_type or "").lower() in CSV_FILE_TYPES or mime_type in {
            "text/csv",
            "application/csv",
        }

    def parse(self, context: ParseContext) -> ParsedDocumentData:
        text = _read_text_file(context.file_path)
        reader = csv.reader(StringIO(text))
        rows = [[_normalize_text(cell) for cell in row] for row in reader]
        rows = [row for row in rows if any(row)]
        table_text = _rows_to_markdown(rows)
        blocks = (
            [
                ParsedBlockData(
                    block_index=0,
                    block_type=BlockType.TABLE,
                    text=table_text,
                    page_number=1,
                    table_metadata={"row_count": len(rows)},
                )
            ]
            if table_text
            else []
        )
        page = ParsedPageData(
            page_number=1,
            logical_page_key="csv",
            text=table_text,
            blocks=blocks,
            parser_metadata={"parser": self.name},
        )
        return self._document(pages=[page], metadata={"source_format": "csv"})


class PlainTextParser(BaseParser):
    name = "plain-text"

    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        return (file_type or "").lower() in TEXT_FILE_TYPES or (mime_type or "").startswith("text/")

    def parse(self, context: ParseContext) -> ParsedDocumentData:
        text = _read_text_file(context.file_path)
        blocks = _blocks_from_plain_text(text)
        page = ParsedPageData(
            page_number=1,
            logical_page_key="text",
            text="\n\n".join(block.text for block in blocks),
            blocks=blocks,
            parser_metadata={"parser": self.name},
        )
        return self._document(pages=[page], metadata={"source_format": "text"})


class ImageOCRParser(BaseParser):
    name = "pytesseract"

    def __init__(self, image_writer: PageImageWriter | None = None) -> None:
        self.image_writer = image_writer or PageImageWriter()

    def supports(self, file_type: str | None, mime_type: str | None) -> bool:
        return (file_type or "").lower() in IMAGE_FILE_TYPES or (mime_type or "").startswith(
            "image/"
        )

    def parse(self, context: ParseContext) -> ParsedDocumentData:
        warnings: list[ParserWarning] = []
        try:
            image = Image.open(context.file_path)
            image.load()
        except (OSError, UnidentifiedImageError) as exc:
            raise DocumentParseError("Image parsing failed.") from exc

        image_key = self.image_writer.write_image(context, image, 1)
        text = ""
        confidence: float | None = None
        ocr_used = False
        if context.settings.ocr_enabled:
            if shutil.which("tesseract") is None:
                warnings.append(
                    ParserWarning(
                        code="ocr_unavailable",
                        message="Tesseract is not installed; image text was not extracted.",
                        scope="page",
                        page_number=1,
                    )
                )
            else:
                data = pytesseract.image_to_data(
                    image,
                    lang=context.settings.ocr_language,
                    output_type=pytesseract.Output.DICT,
                )
                words, confidences = _words_and_confidences(data)
                text = _normalize_text(" ".join(words))
                confidence = _average(confidences)
                ocr_used = True
                if len(text) < context.settings.ocr_min_text_characters:
                    warnings.append(
                        ParserWarning(
                            code="ocr_low_text",
                            message="OCR produced less text than the configured minimum.",
                            scope="page",
                            page_number=1,
                        )
                    )
        else:
            warnings.append(
                ParserWarning(
                    code="ocr_disabled",
                    message="OCR is disabled by configuration.",
                    scope="page",
                    page_number=1,
                )
            )

        blocks = (
            [
                ParsedBlockData(
                    block_index=0,
                    block_type=BlockType.IMAGE_TEXT,
                    text=text,
                    page_number=1,
                    parser_metadata={"ocr_confidence": confidence},
                )
            ]
            if text
            else []
        )
        page = ParsedPageData(
            page_number=1,
            logical_page_key="image",
            text=text,
            blocks=blocks,
            rendered_image_key=image_key,
            width=float(image.width),
            height=float(image.height),
            ocr_used=ocr_used,
            ocr_confidence=confidence,
            warnings=warnings,
            parser_metadata={"parser": self.name},
        )
        return self._document(
            pages=[page],
            warnings=warnings,
            metadata={"source_format": "image"},
            ocr_used=ocr_used,
        )


class ParserRegistry:
    def __init__(self, parsers: Iterable[DocumentParser] | None = None) -> None:
        self.parsers = list(
            parsers
            or [
                PyMuPDFParser(),
                DocxParser(),
                OpenPyXLParser(),
                CSVParser(),
                PlainTextParser(),
                ImageOCRParser(),
            ]
        )

    def get_parser(self, file_type: str | None, mime_type: str | None) -> DocumentParser:
        for parser in self.parsers:
            if parser.supports(file_type, mime_type):
                return parser
        raise UnsupportedParserError()


def _blocks_from_pymupdf(
    block_entries: list[tuple[float, ...]], *, page_number: int
) -> list[ParsedBlockData]:
    blocks: list[ParsedBlockData] = []
    for entry in block_entries:
        if len(entry) < 5:
            continue
        text = _normalize_text(str(entry[4]))
        if not text:
            continue
        block_type = BlockType.HEADING if _looks_like_heading(text) else BlockType.PARAGRAPH
        bbox = {
            "x0": float(entry[0]),
            "y0": float(entry[1]),
            "x1": float(entry[2]),
            "y1": float(entry[3]),
        }
        blocks.append(
            ParsedBlockData(
                block_index=len(blocks),
                block_type=block_type,
                text=text,
                page_number=page_number,
                heading_level=1 if block_type == BlockType.HEADING else None,
                section_path=text if block_type == BlockType.HEADING else None,
                bounding_box=bbox,
            )
        )
    if not blocks:
        return []
    section_path: str | None = None
    for block in blocks:
        if block.block_type == BlockType.HEADING:
            section_path = block.text
        elif section_path:
            block.section_path = section_path
    return blocks


def _blocks_from_plain_text(text: str) -> list[ParsedBlockData]:
    blocks: list[ParsedBlockData] = []
    section_path: str | None = None
    for paragraph in re.split(r"\n\s*\n", text):
        normalized = _normalize_text(paragraph)
        if not normalized:
            continue
        block_type = BlockType.HEADING if _looks_like_heading(normalized) else BlockType.PARAGRAPH
        if block_type == BlockType.HEADING:
            section_path = normalized
        blocks.append(
            ParsedBlockData(
                block_index=len(blocks),
                block_type=block_type,
                text=normalized,
                page_number=1,
                heading_level=1 if block_type == BlockType.HEADING else None,
                section_path=section_path,
            )
        )
    return blocks


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    best = from_bytes(raw).best()
    if best is None:
        return raw.decode("utf-8", errors="replace")
    return str(best)


def _normalize_text(value: str) -> str:
    lines = [" ".join(line.strip().split()) for line in value.replace("\r", "\n").split("\n")]
    return "\n".join(line for line in lines if line).strip()


def _heading_level(style_name: str) -> int | None:
    match = re.search(r"heading\s+([1-6])", style_name)
    if match:
        return int(match.group(1))
    return None


def _looks_like_heading(text: str) -> bool:
    if len(text) > 90 or "\n" in text:
        return False
    if re.match(r"^\d+(\.\d+)*\s+\S+", text):
        return True
    letters = [char for char in text if char.isalpha()]
    return bool(letters) and sum(char.isupper() for char in letters) / len(letters) > 0.7


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    width = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (width - len(row)) for row in rows]
    output = [" | ".join(row).strip() for row in normalized_rows]
    return "\n".join(output)


def _words_and_confidences(data: dict[str, list[object]]) -> tuple[list[str], list[float]]:
    words: list[str] = []
    confidences: list[float] = []
    raw_text = data.get("text", [])
    raw_conf = data.get("conf", [])
    for word, confidence in zip(raw_text, raw_conf, strict=False):
        text = _normalize_text(str(word))
        if not text:
            continue
        words.append(text)
        try:
            parsed_confidence = float(str(confidence))
        except (TypeError, ValueError):
            continue
        if parsed_confidence >= 0:
            confidences.append(parsed_confidence / 100)
    return words, confidences


def _average(values: list[float]) -> float | None:
    if not values:
        return None
    return round(sum(values) / len(values), 4)
